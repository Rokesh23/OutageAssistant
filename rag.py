import os
import glob
import re

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCS_DIR = os.path.join(BASE_DIR, "docs")


def read_docx(file_path):
    """Extracts all text including paragraphs and tables from a .docx file."""
    if not HAS_DOCX:
        return ""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    full_text.append(" | ".join(row_data))
        return "\n".join(full_text)
    except Exception:
        return ""


def load_all_rca_docs():
    """Finds and reads all .docx files in docs/ or root directory."""
    documents = []
    file_paths = []
    
    # Check docs/ directory
    if os.path.exists(DOCS_DIR):
        file_paths.extend(glob.glob(os.path.join(DOCS_DIR, "*.docx")))
        file_paths.extend(glob.glob(os.path.join(DOCS_DIR, "*.txt")))
        file_paths.extend(glob.glob(os.path.join(DOCS_DIR, "*.md")))

    # Check root directory as fallback
    file_paths.extend(glob.glob(os.path.join(BASE_DIR, "*.docx")))

    # Deduplicate and remove temporary Office files (~$)
    unique_paths = list(set(file_paths))

    for path in unique_paths:
        filename = os.path.basename(path)
        if filename.startswith("~$"):
            continue

        if filename.endswith(".docx"):
            content = read_docx(path)
        else:
            try:
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read().strip()
            except Exception:
                content = ""

        if content:
            documents.append({
                "source": filename,
                "text": content
            })

    return documents


def search_documents(query: str, top_k: int = 14):
    documents = load_all_rca_docs()

    if not documents:
        msg = "Sorry, I could not find any RCA documents (.docx) in the repository."
        return msg, [], 0.0

    query_raw = query.strip().lower()

    is_broad_query = any(phrase in query_raw for phrase in [
        "all incidents", "list incidents", "all rca", "show all", "total incidents", "all documents"
    ])

    matches = []
    query_terms = [t for t in re.split(r'\W+', query_raw) if t]

    for item in documents:
        text = item["text"]
        source = item["source"]
        text_lower = text.lower()
        score = 0.0

        if is_broad_query:
            score = 1.0
        else:
            if query_raw in text_lower:
                score += 10.0
            for term in query_terms:
                if len(term) > 2 and term in text_lower:
                    score += 2.0
                elif term in text_lower:
                    score += 0.5

            aliases = {
                "timeout": ["504", "timed out", "slow", "gateway"],
                "504": ["timeout", "gateway", "bad gateway"],
                "502": ["bad gateway", "ui package"],
                "503": ["patching", "unavailable"],
                "ssl": ["cert", "certificate", "err_bad_ssl"]
            }
            for term in query_terms:
                if term in aliases:
                    for alias in aliases[term]:
                        if alias in text_lower:
                            score += 1.5

        if score > 0:
            # Create a compact extract of each doc (first 600 chars) to stay within LLM context limit
            compact_text = text[:600] if len(text) > 600 else text
            formatted_item = {
                "text": text,
                "source": source,
                "score": score
            }
            matches.append((score, formatted_item, f"Source ({source}):\n{compact_text}"))

    matches.sort(key=lambda x: x[0], reverse=True)

    if not matches:
        msg = "Sorry, I could not find any RCA document related to your query in the indexed records."
        return msg, [], 0.0

    selected_matches = matches if is_broad_query else matches[:top_k]
    matched_results = [m[1] for m in selected_matches]
    context_str = "\n\n---\n\n".join([m[2] for m in selected_matches])

    top_score = selected_matches[0][0]
    confidence = min(round((top_score / 10.0) * 100, 1), 100.0)

    return context_str, matched_results, confidence


retrieve = search_documents
